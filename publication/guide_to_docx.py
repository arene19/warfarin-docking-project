#!/usr/bin/env python3
"""Convert PIPELINE_BEGINNER_GUIDE.md to a standalone Word document."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_MD = ROOT / "publication" / "PIPELINE_BEGINNER_GUIDE.md"
DEFAULT_DOCX = ROOT / "PIPELINE_BEGINNER_GUIDE.docx"

ITALIC_TERMS = re.compile(
    r"\b(VKORC1|CYP2C9|HSA|REINVENT4|GROMACS|CHARMM|CGenFF|ParamChem|ChEMBL|RDKit|PyTorch|SMILES|ADMET|IC50|Ki|ROC|AUC|MD|PDBQT|PDB|GPU|CPU|CUDA|MLP|GAT|GCN|RF|ECFP4|QED|TPSA|MW|SA|NLL|DAP|RT)\b"
)
FORMULA_LINE = re.compile(
    r"^(pXC|L_|RMSE|MAE|R²|R\^2|S_total|ΔG|ȳ|ŷ|Σ|max\(|min\(|ρ\s*=)"
)


def add_formatted_text(paragraph, text: str, base_italic: bool = False) -> None:
    text = text.strip()
    if not text:
        return
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if base_italic:
                run.italic = True
        else:
            subparts = re.split(r"(\*[^*]+\*)", part)
            for sub in subparts:
                if not sub:
                    continue
                if sub.startswith("*") and sub.endswith("*") and len(sub) > 2:
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                    continue
                term_parts = ITALIC_TERMS.split(sub)
                for term in term_parts:
                    if not term:
                        continue
                    if ITALIC_TERMS.fullmatch(term):
                        run = paragraph.add_run(term)
                        run.italic = True
                    else:
                        run = paragraph.add_run(term)
                        if base_italic:
                            run.italic = True


def add_heading(doc: Document, text: str, level: int) -> None:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text).strip()
    if text:
        doc.add_heading(text, level=min(level, 3))


def set_repeat_table_header(table) -> None:
    if not table.rows:
        return
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        from docx.oxml import OxmlElement

        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)


def flush_table(doc: Document, table_rows: list[list[str]]) -> None:
    if not table_rows:
        return
    ncols = max(len(row) for row in table_rows)
    normalized = [row + [""] * (ncols - len(row)) for row in table_rows]
    tbl = doc.add_table(rows=len(normalized), cols=ncols)
    tbl.style = "Table Grid"
    for i, row in enumerate(normalized):
        for j, cell in enumerate(row):
            cell_obj = tbl.rows[i].cells[j]
            cell_obj.text = ""
            p = cell_obj.paragraphs[0]
            raw = row[j].strip()
            if i == 0:
                run = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", raw))
                run.bold = True
            else:
                add_formatted_text(p, re.sub(r"\*\*(.+?)\*\*", r"\1", raw))
    set_repeat_table_header(tbl)


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_table = False
    table_rows: list[list[str]] = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("|") and "|" in stripped[1:]:
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue

        if in_table:
            flush_table(doc, table_rows)
            table_rows = []
            in_table = False

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 0)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2)
        elif stripped.startswith("#### "):
            add_heading(doc, stripped[5:], 3)
        elif stripped == "---":
            doc.add_paragraph()
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            add_formatted_text(p, stripped[2:], base_italic=True)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("```"):
            continue
        elif FORMULA_LINE.match(stripped) or stripped.startswith("pXC₅₀"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.italic = True
            run.font.size = Pt(12)
        elif stripped:
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)
        else:
            doc.add_paragraph()

    if in_table:
        flush_table(doc, table_rows)

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


def main() -> None:
    md_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_MD
    docx_path = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_DOCX
    if not md_path.is_file():
        sys.exit(f"Missing {md_path}")
    convert(md_path, docx_path)


if __name__ == "__main__":
    main()
