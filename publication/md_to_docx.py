#!/usr/bin/env python3
"""Convert vkorc1_integrated_workflow_manuscript.md to docx with embedded figures."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "vkorc1_integrated_workflow_manuscript.md"
DOCX_PATH = ROOT / "vkorc1_integrated_workflow_manuscript.docx"

FIG_PATH_RE = re.compile(
    r"^`(publication/output/figures/[^`]+\.png)`\s*$"
)
FIG_CAPTION_RE = re.compile(r"^\*\*Figure\s+(\d+|S\d+)\.\*\*\s*(.*)", re.I)
TABLE_CAPTION_RE = re.compile(r"^\*\*Table\s+(\d+|S\d+)\.\*\*\s*(.*)", re.I)

FIGURE_FILES = {
    "1": "publication/output/figures/figure1_gnn_training_curves.png",
    "2": "publication/output/figures/figure2_baseline_comparison.png",
    "3": "publication/output/figures/figure3_scaffold_split.png",
    "4": "publication/output/figures/figure4_morgan_rf_analysis.png",
    "5": "publication/output/figures/figure5_flexible_redock_spotcheck.png",
    "6": "publication/output/figures/figure6_interaction_heatmap.png",
    "7": "publication/output/figures/figure7_reinvent_vs_training_distribution.png",
    "8": "publication/output/figures/figure8_md_rmsd_rl_gen_37.png",
    "9": "publication/output/figures/figure9_md_hbond_occupancy.png",
}

ITALIC_TERMS = re.compile(
    r"\b(VKORC1|CYP2C9|HSA|REINVENT4|GROMACS|CHARMM|CGenFF|ParamChem|ChEMBL|RDKit|PyTorch)\b"
)


def resolve_figure_path(path_str: str) -> Path | None:
    for path in (ROOT / path_str, ROOT / path_str.replace("publication/", "")):
        if path.is_file():
            return path
    return None


def add_formatted_text(paragraph, text: str, base_italic: bool = False) -> None:
    """Add text with bold (**), italic (*), italic gene names, and Unicode preserved."""
    text = text.strip()
    if not text:
        return
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if base_italic:
                run.italic = True
        else:
            subparts = re.split(r"(\*[^*]+\*)", part)
            for sub in subparts:
                if not sub:
                    continue
                if sub.startswith("*") and sub.endswith("*") and len(sub) > 2:
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                    continue
                term_parts = ITALIC_TERMS.split(sub)
                for term in term_parts:
                    if not term:
                        continue
                    if ITALIC_TERMS.fullmatch(term):
                        run = paragraph.add_run(term)
                        run.italic = True
                    else:
                        run = paragraph.add_run(term)
                        if base_italic:
                            run.italic = True


def add_heading(doc: Document, text: str, level: int) -> None:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text).strip()
    if text:
        doc.add_heading(text, level=min(level, 3))


def embed_figure(doc: Document, path_str: str, caption: str | None = None) -> bool:
    path = resolve_figure_path(path_str)
    if path is None:
        p = doc.add_paragraph()
        run = p.add_run(f"[Figure not found: {path_str}]")
        run.italic = True
        return False

    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(path), width=Inches(6.0))

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_text(cap, caption, base_italic=True)
        for run in cap.runs:
            run.font.size = Pt(10)
    return True


def set_repeat_table_header(table) -> None:
    """Mark first row as repeating on each page (Word)."""
    if not table.rows:
        return
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = trPr.find(qn("w:tblHeader"))
    if tblHeader is None:
        from docx.oxml import OxmlElement

        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)


def flush_table(doc: Document, table_rows: list[list[str]]) -> None:
    if not table_rows:
        return
    tbl = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    tbl.style = "Table Grid"
    for i, row in enumerate(table_rows):
        for j, cell in enumerate(row):
            cell_obj = tbl.rows[i].cells[j]
            cell_obj.text = ""
            p = cell_obj.paragraphs[0]
            raw = row[j].strip()
            if i == 0:
                run = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", raw))
                run.bold = True
            else:
                add_formatted_text(p, re.sub(r"\*\*(.+?)\*\*", r"\1", raw))
    set_repeat_table_header(tbl)


def main() -> None:
    if not MD_PATH.is_file():
        sys.exit(f"Missing {MD_PATH}")

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    in_table = False
    in_references = False
    table_rows: list[list[str]] = []
    seen_figures: set[str] = set()
    pending_fig_caption: str | None = None
    pending_fig_path: str | None = None

    for line in lines:
        stripped = line.strip()

        # Standalone figure path line
        fig_match = FIG_PATH_RE.match(stripped)
        if fig_match:
            path_str = fig_match.group(1)
            if path_str not in seen_figures:
                seen_figures.add(path_str)
                if pending_fig_caption is None:
                    print(f"  [warn] Figure embedded without caption: {path_str}")
                embed_figure(doc, path_str, pending_fig_caption)
                pending_fig_caption = None
                pending_fig_path = None
            continue

        # Figure caption — embed from map if no explicit path follows
        cap_match = FIG_CAPTION_RE.match(stripped)
        if cap_match:
            fig_num = cap_match.group(1)
            pending_fig_caption = f"Figure {fig_num}. {cap_match.group(2).strip()}"
            pending_fig_path = FIGURE_FILES.get(fig_num)
            if pending_fig_path and pending_fig_path not in seen_figures:
                seen_figures.add(pending_fig_path)
                embed_figure(doc, pending_fig_path, pending_fig_caption)
                pending_fig_caption = None
                pending_fig_path = None
            continue

        tbl_cap = TABLE_CAPTION_RE.match(stripped)
        if tbl_cap:
            p = doc.add_paragraph()
            add_formatted_text(
                p, f"Table {tbl_cap.group(1)}. {tbl_cap.group(2).strip()}", base_italic=False
            )
            for run in p.runs:
                run.bold = True
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue

        if in_table:
            flush_table(doc, table_rows)
            table_rows = []
            in_table = False

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 0)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1)
            in_references = stripped[3:].strip() == "References"
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2)
        elif stripped.startswith("#### "):
            add_heading(doc, stripped[5:], 3)
        elif stripped == "---":
            doc.add_paragraph()
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            add_formatted_text(p, stripped[2:], base_italic=True)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            if in_references:
                p = doc.add_paragraph()
                add_formatted_text(p, stripped)
            else:
                p = doc.add_paragraph(style="List Number")
                add_formatted_text(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("pXC₅₀") or stripped.startswith("pXC50"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(p, stripped)
        elif stripped:
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)
        else:
            doc.add_paragraph()

    if in_table:
        flush_table(doc, table_rows)

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH} ({len(seen_figures)} figures embedded)")


if __name__ == "__main__":
    main()
